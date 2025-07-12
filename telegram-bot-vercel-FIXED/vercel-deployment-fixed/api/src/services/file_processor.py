import os
import uuid
import re
from werkzeug.utils import secure_filename
from src.models.bot import db, Document, TextChunk
import PyPDF2
import docx
import markdown
from bs4 import BeautifulSoup

class FileProcessor:
    def __init__(self, upload_folder):
        self.upload_folder = upload_folder
        self.allowed_extensions = {'txt', 'pdf', 'docx', 'md'}
        self.chunk_size = 1000  # Characters per chunk
        self.chunk_overlap = 200  # Overlap between chunks
    
    def is_allowed_file(self, filename):
        """Check if file extension is allowed"""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in self.allowed_extensions
    
    def save_file(self, file, knowledge_base_id):
        """Save uploaded file and create database record"""
        if not file or not self.is_allowed_file(file.filename):
            raise ValueError("Invalid file type")
        
        # Generate unique filename
        original_filename = file.filename
        file_extension = original_filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{uuid.uuid4().hex}.{file_extension}"
        
        # Create knowledge base specific folder
        kb_folder = os.path.join(self.upload_folder, str(knowledge_base_id))
        os.makedirs(kb_folder, exist_ok=True)
        
        # Save file
        file_path = os.path.join(kb_folder, unique_filename)
        file.save(file_path)
        
        # Get file size
        file_size = os.path.getsize(file_path)
        
        # Create database record
        document = Document(
            filename=unique_filename,
            original_filename=original_filename,
            file_path=file_path,
            file_type=file_extension,
            file_size=file_size,
            knowledge_base_id=knowledge_base_id,
            processed=False
        )
        
        db.session.add(document)
        db.session.commit()
        
        return document
    
    def process_document(self, document_id):
        """Process a document and extract text chunks"""
        document = Document.query.get(document_id)
        if not document:
            raise ValueError("Document not found")
        
        if document.processed:
            return True
        
        try:
            # Extract text based on file type
            text_content = self._extract_text(document)
            
            if not text_content:
                raise ValueError("No text content extracted")
            
            # Create text chunks
            chunks = self._create_chunks(text_content)
            
            # Save chunks to database
            for i, chunk_text in enumerate(chunks):
                chunk = TextChunk(
                    content=chunk_text,
                    chunk_index=i,
                    document_id=document.id
                )
                db.session.add(chunk)
            
            # Mark document as processed
            document.processed = True
            db.session.commit()
            
            return True
            
        except Exception as e:
            db.session.rollback()
            raise e
    
    def _extract_text(self, document):
        """Extract text from document based on file type"""
        file_path = document.file_path
        file_type = document.file_type.lower()
        
        if file_type == 'txt':
            return self._extract_text_from_txt(file_path)
        elif file_type == 'pdf':
            return self._extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return self._extract_text_from_docx(file_path)
        elif file_type == 'md':
            return self._extract_text_from_markdown(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_text_from_txt(self, file_path):
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _extract_text_from_pdf(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    def _extract_text_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    def _extract_text_from_markdown(self, file_path):
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to HTML
            html = markdown.markdown(md_content)
            
            # Extract text from HTML
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            
            return text
        except Exception as e:
            raise ValueError(f"Error extracting text from Markdown: {str(e)}")
    
    def _create_chunks(self, text):
        """Split text into overlapping chunks"""
        # Clean text
        text = self._clean_text(text)
        
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # If this is not the last chunk, try to break at a sentence or word boundary
            if end < len(text):
                # Look for sentence boundary (. ! ?)
                sentence_end = text.rfind('.', start, end)
                if sentence_end == -1:
                    sentence_end = text.rfind('!', start, end)
                if sentence_end == -1:
                    sentence_end = text.rfind('?', start, end)
                
                if sentence_end > start:
                    end = sentence_end + 1
                else:
                    # Look for word boundary
                    word_end = text.rfind(' ', start, end)
                    if word_end > start:
                        end = word_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start < 0:
                start = 0
        
        return chunks
    
    def _clean_text(self, text):
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/]', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def delete_document(self, document_id):
        """Delete document and its associated file"""
        document = Document.query.get(document_id)
        if not document:
            return False
        
        try:
            # Delete file from filesystem
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            # Delete from database (chunks will be deleted due to cascade)
            db.session.delete(document)
            db.session.commit()
            
            return True
            
        except Exception as e:
            db.session.rollback()
            raise e
    
    def get_document_info(self, document_id):
        """Get detailed information about a document"""
        document = Document.query.get(document_id)
        if not document:
            return None
        
        chunk_count = len(document.chunks)
        
        return {
            'id': document.id,
            'filename': document.original_filename,
            'file_type': document.file_type,
            'file_size': document.file_size,
            'processed': document.processed,
            'chunk_count': chunk_count,
            'created_at': document.created_at.isoformat() if document.created_at else None
        }

